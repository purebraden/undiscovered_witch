#!/usr/bin/env python3
"""
Script to generate a KDP-ready DOCX file from markdown chapters.

This script:
- Creates a title page
- Creates a copyright page
- Generates a table of contents
- Adds all chapters with proper page breaks
- Ensures each chapter starts on a new page
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_break(doc):
    """Add a page break to the document."""
    doc.add_page_break()


def add_title_page(doc, title, subtitle, author_name="[AUTHOR NAME]"):
    """Add a title page to the document."""
    # Title page section
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add vertical space
    for _ in range(8):
        title_para.add_run().add_break()
    
    # Title
    title_run = title_para.add_run(title.upper())
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    title_para.add_run().add_break()
    title_para.add_run().add_break()
    
    # Subtitle
    if subtitle:
        subtitle_run = title_para.add_run(subtitle)
        subtitle_run.font.size = Pt(24)
        subtitle_run.font.italic = True
        subtitle_run.font.name = 'Times New Roman'
        title_para.add_run().add_break()
        title_para.add_run().add_break()
    
    # Add vertical space
    for _ in range(6):
        title_para.add_run().add_break()
    
    # Author name
    author_run = title_para.add_run(author_name)
    author_run.font.size = Pt(18)
    author_run.font.name = 'Times New Roman'
    
    add_page_break(doc)


def add_copyright_page(doc, title, year="2024", author_name="[AUTHOR NAME]", cover_designer="[COVER DESIGNER]", publisher_name="[PUBLISHER NAME]"):
    """Add a copyright page to the document."""
    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add vertical space
    for _ in range(20):
        copyright_para.add_run().add_break()
    
    # Copyright text
    copyright_text = f"Copyright © {year} by {author_name}\n\n"
    copyright_text += f"All rights reserved.\n\n"
    copyright_text += f"No part of this book may be reproduced or used in any manner "
    copyright_text += f"without the prior written permission of the copyright owner, "
    copyright_text += f"except for the use of brief quotations in a book review.\n\n"
    copyright_text += f"Printed in the United States of America\n\n"
    copyright_text += f"ISBN: [ISBN NUMBER]\n\n"
    copyright_text += f"Cover Design: {cover_designer}\n\n"
    copyright_text += f"Published by {publisher_name}"
    
    copyright_run = copyright_para.add_run(copyright_text)
    copyright_run.font.size = Pt(11)
    copyright_run.font.name = 'Times New Roman'
    
    add_page_break(doc)


def markdown_to_paragraphs(markdown_text, doc):
    """Convert markdown text to Word paragraphs."""
    lines = markdown_text.split('\n')
    
    for line in lines:
        # Handle headers
        if line.startswith('# '):
            # H1 - Chapter title
            para = doc.add_paragraph()
            para.style = 'Heading 1'
            text = line[2:].strip()
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)  # No indent for headings
        
        elif line.startswith('## '):
            # H2
            para = doc.add_paragraph()
            para.style = 'Heading 2'
            text = line[3:].strip()
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
            para.paragraph_format.first_line_indent = Pt(0)  # No indent for headings
        
        elif line.startswith('### '):
            # H3
            para = doc.add_paragraph()
            para.style = 'Heading 3'
            text = line[4:].strip()
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            para.paragraph_format.first_line_indent = Pt(0)  # No indent for headings
        
        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, text, indent=False)
        
        # Handle numbered lists (basic)
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, text, indent=False)
        
        # Handle horizontal rules (just add extra spacing)
        elif line.strip() == '---':
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_after = Pt(12)
            para_format.space_before = Pt(12)
        
        # Empty line
        elif line.strip() == '':
            doc.add_paragraph()
        
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            para.style = 'Normal'
            add_formatted_text(para, line)
    
    return doc


def add_formatted_text(paragraph, text, indent=True):
    """Add text with markdown formatting (bold, italic) to a paragraph."""
    para_format = paragraph.paragraph_format
    para_format.space_after = Pt(6)
    if indent:
        para_format.first_line_indent = Inches(0.5)  # First line indent for paragraphs
    
    # Set default font
    default_font = 'Times New Roman'
    default_size = Pt(12)
    
    # Parse markdown formatting: **bold**, *italic*
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if not part:
            continue
        
        run = paragraph.add_run()
        run.font.name = default_font
        run.font.size = default_size
        
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic
            run.text = part[1:-1]
            run.font.italic = True
        else:
            # Regular text
            run.text = part


def add_table_of_contents(doc):
    """Add a table of contents page."""
    toc_para = doc.add_paragraph()
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_run = toc_para.add_run("TABLE OF CONTENTS")
    toc_run.font.size = Pt(18)
    toc_run.font.bold = True
    toc_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # Add spacing
    
    # Add TOC field (Word will generate the actual TOC)
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Table of Contents"
    
    fldChar2.append(fldChar3)
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar4)
    
    add_page_break(doc)


def process_chapter_file(file_path, doc, is_first_chapter=False):
    """Process a single markdown chapter file and add it to the document."""
    print(f"Processing {file_path.name}...")
    
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Add page break before chapter (except for the very first one)
    if not is_first_chapter:
        add_page_break(doc)
    
    # Process the markdown content
    markdown_to_paragraphs(content, doc)


def get_sorted_chapter_files(book_dir):
    """Get all chapter files sorted in the correct order."""
    book_path = Path(book_dir)
    
    # Define the order
    files_order = []
    
    # Add prologue first
    prologue = book_path / 'prologue.md'
    if prologue.exists():
        files_order.append(prologue)
    
    # Add chapters 1-14
    for i in range(1, 15):
        chapter_file = book_path / f'chapter_{i:02d}.md'
        if chapter_file.exists():
            files_order.append(chapter_file)
    
    # Add epilogue last
    epilogue = book_path / 'epilogue.md'
    if epilogue.exists():
        files_order.append(epilogue)
    
    return files_order


def set_document_styles(doc):
    """Set default document styles for KDP formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15
    paragraph_format.first_line_indent = Inches(0.5)
    
    # Heading 1 style (for chapter titles)
    try:
        heading1 = doc.styles['Heading 1']
        heading1_font = heading1.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_format = heading1.paragraph_format
        heading1_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading1_format.space_after = Pt(12)
        heading1_format.space_before = Pt(12)
    except:
        pass
    
    # Heading 2 style
    try:
        heading2 = doc.styles['Heading 2']
        heading2_font = heading2.font
        heading2_font.name = 'Times New Roman'
        heading2_font.size = Pt(16)
        heading2_font.bold = True
        heading2_format = heading2.paragraph_format
        heading2_format.space_after = Pt(10)
        heading2_format.space_before = Pt(10)
    except:
        pass


def main():
    """Main function to generate the KDP docx file."""
    # Configuration
    book_dir = Path('book1')
    output_file = book_dir / 'Undiscovered_Witch_KDP_Interior_Final.docx'
    
    book_title = "UNDISCOVERED WITCH"
    book_subtitle = "Book One: What Doesn't Settle"
    author_name = "Jessica Braden"
    copyright_year = "2026"
    cover_designer = "Jessica Braden"
    publisher_name = "Jessica Braden"
    
    print(f"Generating KDP-ready DOCX file...")
    print(f"Output: {output_file}")
    
    # Create document
    doc = Document()
    
    # Set document styles
    set_document_styles(doc)
    
    # Add title page
    print("Adding title page...")
    add_title_page(doc, book_title, book_subtitle, author_name)
    
    # Add copyright page
    print("Adding copyright page...")
    add_copyright_page(doc, book_title, copyright_year, author_name, cover_designer, publisher_name)
    
    # Add table of contents
    print("Adding table of contents...")
    add_table_of_contents(doc)
    
    # Process chapters
    chapter_files = get_sorted_chapter_files(book_dir)
    print(f"\nFound {len(chapter_files)} chapter files to process:")
    
    for i, chapter_file in enumerate(chapter_files):
        is_first = (i == 0)
        process_chapter_file(chapter_file, doc, is_first_chapter=is_first)
    
    # Save document
    print(f"\nSaving document to {output_file}...")
    doc.save(str(output_file))
    
    print(f"\nSuccessfully generated {output_file}")
    print("\nNote: You may need to update the TOC in Word:")
    print("  1. Right-click on the TOC and select 'Update Field'")
    print("  2. Or press F9 to update all fields")
    print("\nAlso, please update:")
    print(f"  - Copyright year (currently: {copyright_year})")
    print("  - ISBN number")


if __name__ == '__main__':
    main()

